import docx
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

# try:
#     from deep_translator import GoogleTranslator
#     GOOGLE_READY = True
# except ImportError:
#     GOOGLE_READY = False

# Thêm hỗ trợ Gemini 2.0
try:
    from google import genai
    from google.genai import types
    GEMINI_READY = True
except ImportError:
    GEMINI_READY = False

AI_READY = GEMINI_READY

# LANG_MAP_ISO không còn cần thiết nếu chỉ dùng Gemini

# Fonts hỗ trợ CJK để tránh lỗi ô vuông (tofu)
FONT_MAP = {
    "zh-CN": "SimSun",
    "ja": "MS Gothic",
    "ko": "Malgun Gothic",
    "vi": "Times New Roman",
    "en": "Times New Roman"
}

def apply_font_to_run(run, font_name):
    """Áp dụng mãnh liệt font cho mọi loại ký tự để tránh fallback về font gốc"""
    if not font_name: return
    run.font.name = font_name
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:cs'), font_name)

def translate_blocks_real_ai(texts, target="vi", api_key=None):
    """Dịch thuật thuần túy bằng Gemini 2.0 AI - Không dùng Google Translate API cũ"""
    if not texts: return []
    
    if GEMINI_READY and api_key:
        try:
            client = genai.Client(api_key=api_key)
            combined = "\n---\n".join(texts)
            
            prompt = f"""Bạn là một chuyên gia dịch thuật kỹ thuật đa ngôn ngữ. 
            Hãy dịch các đoạn văn bản sau sang {target}.
            YÊU CẦU:
            1. Giữ nguyên định dạng, không thêm bớt nội dung.
            2. Các đoạn cách nhau bởi dấu '---' xuống dòng.
            3. Trả về đúng số lượng đoạn đã gửi.
            4. Ưu tiên thuật ngữ kỹ thuật chính xác.
            
            VĂN BẢN CẦN DỊCH:
            {combined}"""
            
            response = client.models.generate_content(
                model="gemini-2.0-flash",
                contents=prompt
            )
            
            if response and response.text:
                res = [r.strip() for r in response.text.split("\n---\n")]
                if len(res) == len(texts): return res
                # Fallback nếu split lỗi
                return [r.strip() for r in response.text.split("---") if r.strip()][:len(texts)]
        except Exception as e:
            print(f"Gemini Error: {e}")

    # Fallback cuối cùng nếu không có API Key hoặc lỗi
    return [f"[Yêu cầu API Key để dịch sang {target}] {txt}" for txt in texts]

def translate_docx_v823(input_p, output_p, target_l="vi", is_bi=False):
    """Quy trình Dịch thuật Master V823 (Fix lỗi giãn chữ Justify)"""
    try:
        doc = Document(input_p)
        orig_paras = list(doc.paragraphs)
        
        for p in orig_paras:
            if not p.text.strip(): continue
            
            # Dịch từng đoạn đơn lẻ để giữ style (Hoặc có thể gom block nếu muốn tốc độ)
            # Ở đây dùng tham số api_key để truyền vào engine
            trans_texts = translate_blocks_real_ai([p.text], target_l, api_key=os.getenv("GOOGLE_API_KEY"))
            if not trans_texts: continue
            trans_text = trans_texts[0]
            
            if is_bi:
                # Tạo Paragraph mới ngay dưới bản gốc bằng insert_paragraph_after-trick
                new_p = p.insert_paragraph_before("") 
                run = new_p.add_run(f"({trans_text})")
                run.italic = True
                run.font.size = Pt(11)
                
                # Áp dụng font CJK nếu cần
                target_font = FONT_MAP.get(target_l, "Arial")
                apply_font_to_run(run, target_font)
                
                new_p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                
                # Hoán đổi element để Gốc nằm trên - Dịch nằm dưới
                p_elt = p._element
                new_p_elt = new_p._element
                p_elt.addnext(new_p_elt)
            else:
                p.text = trans_text
                # Với trường hợp thay thế hoàn toàn, cũng cần set font cho tất cả runs
                target_font = FONT_MAP.get(target_l, "Arial")
                for run in p.runs:
                    apply_font_to_run(run, target_font)

        # Xử lý bảng biểu (Table)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        t_cell_list = translate_blocks_real_ai([cell.text], target_l, api_key=os.getenv("GOOGLE_API_KEY"))
                        if t_cell_list:
                            t_cell = t_cell_list[0]
                            if is_bi:
                                new_para = cell.add_paragraph(f"({t_cell})")
                                run = new_para.runs[0]
                                run.italic = True
                                
                                # Áp dụng font CJK nếu cần
                                target_font = FONT_MAP.get(target_l, "Arial")
                                apply_font_to_run(run, target_font)
                                
                                new_para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                            else:
                                cell.text = t_cell
                                target_font = FONT_MAP.get(target_l, "Arial")
                                for p_cell in cell.paragraphs:
                                    for run in p_cell.runs:
                                        apply_font_to_run(run, target_font)
                            
        doc.save(output_p)
        return True
    except Exception as e:
        print(f"Lỗi V823: {e}")
        return False

def translate_docx(i, o, lang="vi", bi=False, api_key=None):
    # Set env var tạm thời để engine bên trên có thể dùng (hoặc truyền tham số)
    if api_key: os.environ["GOOGLE_API_KEY"] = api_key
    return translate_docx_v823(i, o, target_l=lang, is_bi=bi)
